import docx
import sys
import os

doc_path = "/home/mehdi/VANET_Project/Docker_files/runs_roc_pr_comparison.docx"
if not os.path.exists(doc_path):
    print(f"Error: {doc_path} does not exist.")
    sys.exit(1)

doc = docx.Document(doc_path)
print(f"Loaded {doc_path}. It has {len(doc.tables)} tables.")

for idx, table in enumerate(doc.tables):
    if len(table.rows) > 0:
        first_row_text = [cell.text.strip().replace("\n", " ") for cell in table.rows[0].cells]
        print(f"Table {idx}: cols={len(table.columns)}, rows={len(table.rows)} | First row: {first_row_text[:4]}")
        has_keywords = False
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.lower()
                if "reproducibility" in txt or "map source" in txt or "matched-seed" in txt:
                    has_keywords = True
                    break
        if has_keywords:
            print(f"  -> Found keywords in Table {idx}!")
            for r_idx, row in enumerate(table.rows):
                row_text = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                print(f"     Row {r_idx}: {row_text}")
